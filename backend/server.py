"""
Atlas AI - Enterprise Multimodal Agentic RAG Platform
Backend API Server
"""

from fastapi import FastAPI, APIRouter, UploadFile, File, BackgroundTasks, HTTPException
from starlette.responses import StreamingResponse
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from typing import List, Optional, TypedDict
from datetime import datetime, timezone
from pathlib import Path
from google import genai
from google.genai import types
import os
import uuid
import json
import logging
import re
import shutil
import numpy as np

# Vector database
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct, PointIdsList, Filter, FieldCondition, MatchAny
from sklearn.feature_extraction.text import HashingVectorizer

# Document parsers
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image
import pytesseract




# ============================================================
# CONFIGURATION
# ============================================================

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

UPLOAD_DIR = ROOT_DIR / 'uploads'
UPLOAD_DIR.mkdir(exist_ok=True)

MONGO_URL = os.environ['MONGO_URL']
DB_NAME = os.environ['DB_NAME']
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
gemini_client = genai.Client(api_key=GEMINI_API_KEY)

client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]

# Qdrant vector database (in-memory)
vector_store = QdrantClient(":memory:")
QDRANT_COLLECTION = "atlas_chunks"
EMBEDDING_DIM = 384

# Embedding service (stateless HashingVectorizer - consistent dimensions, no model download)
_vectorizer = HashingVectorizer(n_features=EMBEDDING_DIM, alternate_sign=False, norm='l2')


def embed_text(text: str) -> list:
    """Generate embedding vector for a single text."""
    return _vectorizer.transform([text]).toarray()[0].tolist()


def embed_texts(texts: list) -> list:
    """Generate embedding vectors for multiple texts."""
    return _vectorizer.transform(texts).toarray().tolist()


logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("atlas-ai")

app = FastAPI(title="Atlas AI", version="2.0.0")
api_router = APIRouter(prefix="/api")


# ============================================================
# PYDANTIC MODELS
# ============================================================

class DocumentOut(BaseModel):
    id: str
    filename: str
    file_type: str
    file_size: int
    status: str
    chunks_count: int
    selected: bool = True
    error_message: Optional[str] = None
    created_at: str
    updated_at: str


class ConversationOut(BaseModel):
    id: str
    title: str
    message_count: int
    created_at: str
    updated_at: str


class MessageOut(BaseModel):
    id: str
    role: str
    content: str
    sources: list = []
    created_at: str


class ConversationDetailOut(BaseModel):
    id: str
    title: str
    messages: List[MessageOut]
    created_at: str
    updated_at: str


class ChatRequest(BaseModel):
    question: str
    conversation_id: Optional[str] = None
    selected_document_ids: Optional[List[str]] = None


class BulkDeleteRequest(BaseModel):
    document_ids: List[str]


class SelectDocumentsRequest(BaseModel):
    document_ids: List[str]
    selected: bool


class HealthOut(BaseModel):
    status: str
    version: str
    documents_count: int
    chunks_count: int


# ============================================================
# DOCUMENT PARSERS
# ============================================================

def parse_pdf(file_path: str) -> list:
    """Extract text from PDF using PyMuPDF. Returns list of (page_num, text)."""
    doc = fitz.open(file_path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text().strip()
        if text:
            pages.append((i + 1, text))
    doc.close()
    return pages


def parse_docx(file_path: str) -> list:
    """Extract text from DOCX. Returns list of (page_num, text)."""
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = '\n\n'.join(paragraphs)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                full_text += '\n' + ' | '.join(cells)

    return [(1, full_text)] if full_text.strip() else []


def parse_image_ocr(file_path: str) -> str:
    """Extract text from image using Tesseract OCR."""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        logger.warning(f"Tesseract OCR failed: {e}")
        return ""


import base64

async def analyze_image_with_vision(image_path: str, prompt: str) -> str:
    with open(image_path, "rb") as f:
        image_bytes = f.read()

    response = gemini_client.models.generate_content(
        model="gemini-2.5-flash",
        contents=[
            prompt,
            types.Part.from_bytes(
                data=image_bytes,
                mime_type="image/png"
            )
        ]
    )

    return response.text

# ============================================================
# CHUNKING SERVICE
# ============================================================

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150) -> list:
    """Split text into overlapping chunks with sentence boundary awareness."""
    if not text or not text.strip():
        return []

    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + chunk_size, text_len)

        if end < text_len:
            boundary = max(
                text.rfind('. ', start, end),
                text.rfind('\n', start, end),
                text.rfind('! ', start, end),
                text.rfind('? ', start, end),
            )
            if boundary > start + chunk_size // 3:
                end = boundary + 1

        chunk = text[start:end].strip()
        if chunk and len(chunk) > 20:
            chunks.append(chunk)

        start = end - overlap if end < text_len else text_len

    return chunks


# ============================================================
# DOCUMENT PROCESSING PIPELINE
# ============================================================

async def process_document(doc_id: str, file_path: str, file_type: str, mime_type: str, filename: str):
    """Full document processing pipeline: parse -> chunk -> store."""
    try:
        logger.info(f"Processing document {doc_id} ({file_type})")
        pages = []

        if file_type == 'pdf':
            pages = parse_pdf(file_path)
        elif file_type in ('docx', 'doc'):
            pages = parse_docx(file_path)
        elif file_type in ('png', 'jpg', 'jpeg', 'webp', 'gif', 'bmp'):
            ocr_text = parse_image_ocr(file_path)
            vision_text = await analyze_image_with_vision(file_path, mime_type)

            combined_parts = []
            if ocr_text:
                combined_parts.append(f"OCR Extracted Text:\n{ocr_text}")
            if vision_text:
                combined_parts.append(f"Image Analysis:\n{vision_text}")

            combined = '\n\n'.join(combined_parts) if combined_parts else "No text could be extracted from this image."
            pages = [(1, combined)]
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Chunk all pages
        all_chunks = []
        for page_num, page_text in pages:
            text_chunks = chunk_text(page_text)
            for i, chunk in enumerate(text_chunks):
                all_chunks.append({
                    "id": str(uuid.uuid4()),
                    "document_id": doc_id,
                    "text": chunk,
                    "metadata": {
                        "filename": filename,
                        "page": page_num,
                        "chunk_index": i,
                        "file_type": file_type
                    },
                    "created_at": datetime.now(timezone.utc).isoformat()
                })

        if all_chunks:
            await db.chunks.insert_many(all_chunks)

            # Generate embeddings and store in Qdrant
            try:
                texts = [c["text"] for c in all_chunks]
                vectors = embed_texts(texts)
                points = [
                    PointStruct(
                        id=c["id"],
                        vector=vec,
                        payload={
                            "document_id": doc_id,
                            "text": c["text"],
                            "metadata": c["metadata"],
                        }
                    )
                    for c, vec in zip(all_chunks, vectors)
                ]
                vector_store.upsert(collection_name=QDRANT_COLLECTION, points=points)
                logger.info(f"Stored {len(points)} vectors in Qdrant for doc {doc_id}")
            except Exception as e:
                logger.warning(f"Qdrant upsert in process_document failed: {e}")

        await db.documents.update_one(
            {"id": doc_id},
            {"$set": {
                "status": "ready",
                "chunks_count": len(all_chunks),
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        logger.info(f"Document {doc_id} processed: {len(all_chunks)} chunks created")

    except Exception as e:
        logger.error(f"Error processing document {doc_id}: {e}")
        await db.documents.update_one(
            {"id": doc_id},
            {"$set": {
                "status": "error",
                "error_message": str(e),
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )


# ============================================================
# RETRIEVAL SERVICE (Hybrid: Qdrant Dense + MongoDB BM25 + RRF)
# ============================================================

def reciprocal_rank_fusion(results_lists: list, k: int = 60) -> list:
    """Combine multiple ranked result lists using Reciprocal Rank Fusion."""
    scores = {}
    doc_map = {}

    for results in results_lists:
        for rank, result in enumerate(results):
            rid = result.get("id", str(rank))
            scores[rid] = scores.get(rid, 0) + 1.0 / (k + rank + 1)
            if rid not in doc_map:
                doc_map[rid] = result

    sorted_ids = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    return [
        {**doc_map[rid], "rrf_score": round(score, 4)}
        for rid, score in sorted_ids
        if rid in doc_map
    ]


async def hybrid_retrieval(query: str, search_query: str = "", top_k: int = 8, document_ids: list = None) -> list:
    """Hybrid retrieval: Qdrant dense + MongoDB BM25 + RRF reranking. Optionally filtered by document_ids."""
    effective_query = search_query or query

    # Build Qdrant filter for document selection
    qdrant_filter = None
    if document_ids:
        qdrant_filter = Filter(
            must=[FieldCondition(key="document_id", match=MatchAny(any=document_ids))]
        )

    # 1. Qdrant dense search
    qdrant_results = []
    try:
        query_vector = embed_text(effective_query)
        result = vector_store.query_points(
            collection_name=QDRANT_COLLECTION,
            query=query_vector,
            query_filter=qdrant_filter,
            limit=top_k,
        )
        for hit in result.points:
            payload = hit.payload or {}
            qdrant_results.append({
                "id": str(hit.id),
                "text": payload.get("text", ""),
                "metadata": payload.get("metadata", {}),
                "score": round(hit.score, 4),
                "retrieval_method": "dense",
            })
    except Exception as e:
        logger.warning(f"Qdrant dense search failed: {e}")

    # 2. MongoDB BM25 text search
    mongo_results = []
    try:
        mongo_query = {"$text": {"$search": effective_query}}
        if document_ids:
            mongo_query["document_id"] = {"$in": document_ids}
        cursor = db.chunks.find(
            mongo_query,
            {"score": {"$meta": "textScore"}, "_id": 0}
        ).sort([("score", {"$meta": "textScore"})]).limit(top_k)
        raw = await cursor.to_list(top_k)
        for r in raw:
            mongo_results.append({
                "id": r.get("id", ""),
                "text": r.get("text", ""),
                "metadata": r.get("metadata", {}),
                "score": round(r.get("score", 0), 4),
                "retrieval_method": "bm25",
            })
    except Exception as e:
        logger.warning(f"MongoDB text search failed: {e}")

    # 3. Fallback regex if both empty
    if not qdrant_results and not mongo_results:
        terms = [t for t in query.split() if len(t) > 2][:8]
        if terms:
            pattern = '|'.join(re.escape(t) for t in terms)
            try:
                regex_query = {"text": {"$regex": pattern, "$options": "i"}}
                if document_ids:
                    regex_query["document_id"] = {"$in": document_ids}
                cursor = db.chunks.find(regex_query, {"_id": 0}).limit(top_k)
                raw = await cursor.to_list(top_k)
                return [{
                    "id": r.get("id", ""), "text": r.get("text", ""),
                    "metadata": r.get("metadata", {}), "score": 0,
                    "retrieval_method": "regex"
                } for r in raw]
            except Exception:
                pass
        return []

    # 4. Reciprocal Rank Fusion
    all_lists = [l for l in [qdrant_results, mongo_results] if l]
    if len(all_lists) == 1:
        return all_lists[0][:top_k]
    return reciprocal_rank_fusion(all_lists)[:top_k]


def build_context(chunks: list) -> str:
    """Build context string from retrieved chunks for the reasoning agent."""
    if not chunks:
        return "No relevant documents found in the knowledge base."

    parts = []
    for i, chunk in enumerate(chunks):
        meta = chunk.get("metadata", {})
        filename = meta.get("filename", "Unknown")
        page = meta.get("page", "?")
        method = chunk.get("retrieval_method", "hybrid")
        parts.append(f"[Source {i + 1}: {filename}, Page {page} | Retrieved via {method}]\n{chunk['text']}")

    return "\n\n---\n\n".join(parts)


# ============================================================
# AGENT PIPELINE (Planner -> Retrieval -> Memory -> Reasoning)
# ============================================================

async def call_llm(system_message: str, user_text: str) -> str:
    """Call Gemini and return the full response."""

    prompt = f"""
System Instructions:
{system_message}

User:
{user_text}
"""

    response = gemini_client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
    )

    return response.text


async def planner_agent(question: str) -> dict:
    """Planner Agent: classify query intent and extract optimized search parameters."""
    plan_prompt = """You are a query planning agent for a document RAG system.
Analyze the user's question and output ONLY a JSON object:
{"intent": "summarize|extract|compare|explain|search|general", "search_query": "optimized keywords for document search", "entities": ["key", "terms"]}
Output ONLY valid JSON."""

    try:
        raw = await call_llm(plan_prompt, question)
        # Parse JSON from response - handle multi-line and nested objects
        brace_depth = 0
        start_idx = -1
        for i, ch in enumerate(raw):
            if ch == '{':
                if brace_depth == 0:
                    start_idx = i
                brace_depth += 1
            elif ch == '}':
                brace_depth -= 1
                if brace_depth == 0 and start_idx >= 0:
                    return json.loads(raw[start_idx:i + 1])
    except Exception as e:
        logger.warning(f"Planner agent failed: {e}")

    return {"intent": "general", "search_query": question, "entities": question.split()[:5]}


async def memory_agent(messages: list) -> str:
    """Memory Agent: summarize long conversations, keep recent messages intact."""
    if not messages:
        return ""

    if len(messages) <= 6:
        parts = []
        for m in messages:
            role = "User" if m["role"] == "user" else "Atlas AI"
            parts.append(f"{role}: {m['content'][:500]}")
        return "\n".join(parts)

    # Summarize older messages, keep last 4 in full
    older = messages[:-4]
    recent = messages[-4:]

    older_text = "\n".join([
        f"{'User' if m['role'] == 'user' else 'AI'}: {m['content'][:300]}"
        for m in older
    ])

    try:
        summary = await call_llm(
            "Summarize this conversation in 2-3 concise sentences. Focus on key questions and findings.",
            f"Conversation:\n{older_text}"
        )
    except Exception:
        summary = "Previous conversation context available."

    recent_text = "\n".join([
        f"{'User' if m['role'] == 'user' else 'Atlas AI'}: {m['content'][:500]}"
        for m in recent
    ])

    return f"CONVERSATION SUMMARY:\n{summary}\n\nRECENT MESSAGES:\n{recent_text}"


# ============================================================
# LANGGRAPH ORCHESTRATION — StateGraph, Nodes, Edges, Compile
# ============================================================

from langgraph.graph import StateGraph, END, START


class AgentState(TypedDict, total=False):
    """Typed shared state flowing through the LangGraph agent pipeline."""
    # — Input (set before graph invocation)
    user_query: str
    conversation_id: str
    selected_document_ids: list
    conversation_messages: list
    # — Intermediate (populated by nodes)
    plan: dict
    document_ids: list
    needs_vision: bool
    retrieved_chunks: list
    sources: list
    conversation_summary: str
    reasoning_context: str
    # — Output
    citations: list
    final_response: str


# ---- Node definitions ---------------------------------------------------- #

async def planner_node(state: AgentState) -> dict:
    """PlannerAgent — classify query intent, extract optimised search terms."""
    plan = await planner_agent(state["user_query"])
    logger.info(
        f"[LangGraph] PlannerAgent: intent={plan.get('intent')}, "
        f"search_query={plan.get('search_query', '')[:60]}"
    )
    return {"plan": plan}


async def document_node(state: AgentState) -> dict:
    """DocumentAgent — resolve selected document IDs for retrieval filtering."""
    raw = state.get("selected_document_ids") or []
    doc_ids = raw if raw else None
    return {"document_ids": doc_ids}


async def vision_node(state: AgentState) -> dict:
    """VisionAgent — flag whether the query requires visual understanding."""
    vision_kw = {
        "image", "diagram", "picture", "photo", "screenshot",
        "figure", "chart", "graph", "visual", "illustration",
    }
    needs = any(kw in state["user_query"].lower() for kw in vision_kw)
    return {"needs_vision": needs}


async def retrieval_node(state: AgentState) -> dict:
    """RetrievalAgent — hybrid search (Qdrant dense + MongoDB BM25 + RRF)."""
    plan = state.get("plan", {})
    search_q = plan.get("search_query", state["user_query"])
    doc_ids = state.get("document_ids")

    chunks = await hybrid_retrieval(
        state["user_query"], search_query=search_q, document_ids=doc_ids
    )

    sources = []
    for i, c in enumerate(chunks):
        meta = c.get("metadata", {})
        sources.append({
            "index": i + 1,
            "id": c.get("id", ""),
            "text": c["text"][:300],
            "full_text": c["text"],
            "filename": meta.get("filename", "Unknown"),
            "page": meta.get("page", 0),
            "file_type": meta.get("file_type", ""),
            "score": c.get("rrf_score", c.get("score", 0)),
            "retrieval_method": c.get("retrieval_method", "hybrid"),
        })

    return {"retrieved_chunks": chunks, "sources": sources}


async def memory_node(state: AgentState) -> dict:
    """MemoryAgent — summarise long conversations for context-window efficiency."""
    messages = state.get("conversation_messages", [])
    summary = await memory_agent(messages)
    return {"conversation_summary": summary}


async def reasoning_node(state: AgentState) -> dict:
    """ReasoningAgent — assemble the full system prompt / reasoning context.

    The actual LLM streaming call is made *outside* the graph so that SSE
    tokens flow to the client in real time.  This node only prepares the
    prompt that will drive that call.
    """
    context_text = build_context(state.get("retrieved_chunks", []))
    plan = state.get("plan", {})
    intent = plan.get("intent", "general")
    memory = state.get("conversation_summary", "")

    task_map = {
        "summarize": "Provide a comprehensive summary covering all key points from the documents.",
        "compare": "Compare and contrast information from different sources systematically.",
        "extract": "Extract the specific information requested with precise details and data.",
        "explain": "Explain the concept, diagram, or architecture clearly with supporting context.",
        "search": "Find and present the most relevant information with full context.",
        "general": "Provide a thorough, well-structured answer based on the documents.",
    }

    vision_note = (
        "\nNote: This query may involve visual content. "
        "Reference any image-analysis results in the context."
        if state.get("needs_vision") else ""
    )

    system_prompt = (
        "You are Atlas AI, an enterprise document analysis assistant powered by "
        "a LangGraph agentic RAG pipeline with hybrid retrieval "
        "(Qdrant dense search + MongoDB BM25 + Reciprocal Rank Fusion).\n\n"
        f"QUERY INTENT: {intent}\n"
        f"TASK: {task_map.get(intent, task_map['general'])}{vision_note}\n\n"
        f"RETRIEVED DOCUMENT CONTEXT:\n{context_text}\n\n"
        + (f"CONVERSATION MEMORY:\n{memory}\n\n" if memory else "")
        + "CITATION RULES:\n"
        "- ALWAYS cite sources using [Source N] notation (e.g., [Source 1], [Source 2])\n"
        "- Every factual claim MUST reference its source\n"
        "- If context is insufficient, explicitly state what information is missing\n"
        "- Use Markdown formatting: headers, lists, bold, code blocks\n"
        "- Be precise, thorough, and professional"
    )

    return {"reasoning_context": system_prompt}


async def citation_node(state: AgentState) -> dict:
    """CitationAgent — extract and structure citation metadata from sources."""
    sources = state.get("sources", [])
    citations = [
        {
            "index": s["index"],
            "filename": s["filename"],
            "page": s["page"],
            "file_type": s.get("file_type", ""),
        }
        for s in sources
    ]
    return {"citations": citations}


# ---- Graph wiring -------------------------------------------------------- #

_graph_builder = StateGraph(AgentState)

_graph_builder.add_node("planner_agent", planner_node)
_graph_builder.add_node("document_agent", document_node)
_graph_builder.add_node("vision_agent", vision_node)
_graph_builder.add_node("retrieval_agent", retrieval_node)
_graph_builder.add_node("memory_agent", memory_node)
_graph_builder.add_node("reasoning_agent", reasoning_node)
_graph_builder.add_node("citation_agent", citation_node)

_graph_builder.add_edge(START, "planner_agent")
_graph_builder.add_edge("planner_agent", "document_agent")
_graph_builder.add_edge("document_agent", "vision_agent")
_graph_builder.add_edge("vision_agent", "retrieval_agent")
_graph_builder.add_edge("retrieval_agent", "memory_agent")
_graph_builder.add_edge("memory_agent", "reasoning_agent")
_graph_builder.add_edge("reasoning_agent", "citation_agent")
_graph_builder.add_edge("citation_agent", END)

# Compile once at module load — reused for every /chat request
atlas_graph = _graph_builder.compile()


# ============================================================
# API ROUTES
# ============================================================

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'png', 'jpg', 'jpeg', 'webp', 'gif', 'bmp'}


@api_router.post("/upload", response_model=DocumentOut)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
):
    """Upload and process a document (PDF, DOCX, or Image)."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {ext}. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    doc_id = str(uuid.uuid4())
    file_dir = UPLOAD_DIR / doc_id
    file_dir.mkdir(parents=True, exist_ok=True)
    file_path = file_dir / file.filename

    content = await file.read()
    with open(file_path, 'wb') as f:
        f.write(content)

    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "id": doc_id,
        "filename": file.filename,
        "file_type": ext,
        "file_size": len(content),
        "file_path": str(file_path),
        "status": "processing",
        "chunks_count": 0,
        "selected": True,
        "error_message": None,
        "created_at": now,
        "updated_at": now,
    }
    await db.documents.insert_one(doc)

    background_tasks.add_task(
        process_document, doc_id, str(file_path), ext,
        file.content_type or "application/octet-stream", file.filename
    )

    logger.info(f"Scheduled background task for {doc_id}")


    doc.pop("_id", None)
    doc.pop("file_path", None)
    return doc


@api_router.get("/documents", response_model=List[DocumentOut])
async def list_documents():
    """List all uploaded documents."""
    docs = await db.documents.find(
        {}, {"_id": 0, "file_path": 0}
    ).sort("created_at", -1).to_list(100)
    return docs


@api_router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document and its chunks."""
    doc = await db.documents.find_one({"id": doc_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Delete from Qdrant vector store
    try:
        chunk_docs = await db.chunks.find(
            {"document_id": doc_id}, {"id": 1, "_id": 0}
        ).to_list(10000)
        if chunk_docs:
            ids_to_delete = [c["id"] for c in chunk_docs]
            vector_store.delete(
                collection_name=QDRANT_COLLECTION,
                points_selector=PointIdsList(points=ids_to_delete),
            )
    except Exception as e:
        logger.warning(f"Qdrant deletion warning: {e}")

    await db.chunks.delete_many({"document_id": doc_id})
    await db.documents.delete_one({"id": doc_id})

    file_dir = UPLOAD_DIR / doc_id
    if file_dir.exists():
        shutil.rmtree(file_dir, ignore_errors=True)

    return {"status": "deleted", "id": doc_id}


@api_router.post("/documents/delete-bulk")
async def bulk_delete_documents(request: BulkDeleteRequest):
    """Delete multiple documents and their chunks."""
    deleted = []
    for doc_id in request.document_ids:
        doc = await db.documents.find_one({"id": doc_id})
        if not doc:
            continue
        # Delete from Qdrant
        try:
            chunk_docs = await db.chunks.find(
                {"document_id": doc_id}, {"id": 1, "_id": 0}
            ).to_list(10000)
            if chunk_docs:
                ids_to_delete = [c["id"] for c in chunk_docs]
                vector_store.delete(
                    collection_name=QDRANT_COLLECTION,
                    points_selector=PointIdsList(points=ids_to_delete),
                )
        except Exception as e:
            logger.warning(f"Qdrant bulk deletion warning for {doc_id}: {e}")
        await db.chunks.delete_many({"document_id": doc_id})
        await db.documents.delete_one({"id": doc_id})
        file_dir = UPLOAD_DIR / doc_id
        if file_dir.exists():
            shutil.rmtree(file_dir, ignore_errors=True)
        deleted.append(doc_id)
    return {"status": "deleted", "deleted_ids": deleted, "count": len(deleted)}


@api_router.patch("/documents/select")
async def select_documents(request: SelectDocumentsRequest):
    """Update selection state for multiple documents."""
    result = await db.documents.update_many(
        {"id": {"$in": request.document_ids}},
        {"$set": {"selected": request.selected, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    return {"status": "updated", "modified_count": result.modified_count}


@api_router.post("/chat")
async def chat_endpoint(request: ChatRequest):
    """Chat with documents via the compiled LangGraph agent pipeline.

    Graph: PlannerAgent → DocumentAgent → VisionAgent → RetrievalAgent
           → MemoryAgent → ReasoningAgent → CitationAgent
    Streaming: the LLM call is made *after* the graph returns so that SSE
    tokens flow to the client in real time.
    """

    async def generate():
        try:
            conv_id = request.conversation_id or str(uuid.uuid4())
            now = datetime.now(timezone.utc).isoformat()

            # Get or create conversation
            conv = await db.conversations.find_one({"id": conv_id})
            if not conv:
                conv = {
                    "id": conv_id,
                    "title": request.question[:60].strip(),
                    "messages": [],
                    "created_at": now,
                    "updated_at": now,
                }
                await db.conversations.insert_one(conv)

            # Save user message
            user_msg = {
                "id": str(uuid.uuid4()),
                "role": "user",
                "content": request.question,
                "sources": [],
                "created_at": now,
            }
            await db.conversations.update_one(
                {"id": conv_id},
                {"$push": {"messages": user_msg}, "$set": {"updated_at": now}}
            )

            # ===== INVOKE COMPILED LANGGRAPH =====
            initial_state: AgentState = {
                "user_query": request.question,
                "conversation_id": conv_id,
                "selected_document_ids": request.selected_document_ids or [],
                "conversation_messages": conv.get("messages", []),
            }

            final_state = await atlas_graph.ainvoke(initial_state)

            # Read graph outputs
            sources = final_state.get("sources", [])
            plan = final_state.get("plan", {})
            citations = final_state.get("citations", [])
            system_prompt = final_state.get("reasoning_context", "You are Atlas AI.")

            # Send sources to the client
            yield f"event: sources\ndata: {json.dumps({'sources': sources, 'conversation_id': conv_id})}\n\n"

            # Stream LLM response using the reasoning context prepared by the graph
            # Stream response using Gemini

                    # Stream response using Gemini
            full_response = ""

            prompt = f"""
System Instructions:
{system_prompt}

User:
{request.question}
"""

            response = gemini_client.models.generate_content_stream(
                model="gemini-2.5-flash",
                contents=prompt,
            )

            for chunk in response:
                if hasattr(chunk, "text") and chunk.text:
                    full_response += chunk.text
                    yield (
                        f"event: token\n"
                        f"data: {json.dumps({'content': chunk.text})}\n\n"
                    )

            # Persist assistant message with citation metadata
            ai_msg = {
                "id": str(uuid.uuid4()),
                "role": "assistant",
                "content": full_response,
                "sources": citations,
                "created_at": datetime.now(timezone.utc).isoformat(),
            }

            await db.conversations.update_one(
                {"id": conv_id},
                {
                    "$push": {"messages": ai_msg},
                    "$set": {
                        "updated_at": datetime.now(timezone.utc).isoformat()
                    },
                },
            )

            yield (
                f"event: done\ndata: "
                f"{json.dumps({'conversation_id': conv_id, 'message_id': ai_msg['id'], 'plan': plan})}\n\n"
            )

        except Exception as e:
            logger.error(f"Chat error: {e}", exc_info=True)
            yield f"event: error\ndata: {json.dumps({'message': str(e)})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


@api_router.get("/conversations", response_model=List[ConversationOut])
async def list_conversations():
    """List all conversations."""
    convs = await db.conversations.find(
        {}, {"_id": 0}
    ).sort("updated_at", -1).to_list(50)

    result = []
    for c in convs:
        result.append({
            "id": c["id"],
            "title": c.get("title", "Untitled"),
            "message_count": len(c.get("messages", [])),
            "created_at": c["created_at"],
            "updated_at": c["updated_at"],
        })
    return result


@api_router.get("/conversations/{conv_id}", response_model=ConversationDetailOut)
async def get_conversation(conv_id: str):
    """Get a conversation with all messages."""
    conv = await db.conversations.find_one({"id": conv_id}, {"_id": 0})
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return conv


@api_router.delete("/conversations/{conv_id}")
async def delete_conversation(conv_id: str):
    """Delete a conversation."""
    result = await db.conversations.delete_one({"id": conv_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return {"status": "deleted", "id": conv_id}


@api_router.get("/health", response_model=HealthOut)
async def health_check():
    """Health check endpoint."""
    docs_count = await db.documents.count_documents({})
    chunks_count = await db.chunks.count_documents({})
    return {
        "status": "healthy",
        "version": "2.0.0",
        "documents_count": docs_count,
        "chunks_count": chunks_count,
    }


# ============================================================
# APP SETUP
# ============================================================

app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def startup():
    """Create indexes, Qdrant collection, and re-index existing chunks."""
    UPLOAD_DIR.mkdir(exist_ok=True)

    # MongoDB indexes
    try:
        await db.chunks.create_index([("text", "text")])
        await db.chunks.create_index("document_id")
        await db.documents.create_index("id", unique=True)
        await db.conversations.create_index("id", unique=True)
        logger.info("MongoDB indexes created")
    except Exception as e:
        logger.warning(f"Index creation warning: {e}")

    # Qdrant collection
    try:
        collections = vector_store.get_collections().collections
        exists = any(c.name == QDRANT_COLLECTION for c in collections)
        if not exists:
            vector_store.create_collection(
                collection_name=QDRANT_COLLECTION,
                vectors_config=VectorParams(
                    size=EMBEDDING_DIM,
                    distance=Distance.COSINE,
                ),
            )
            logger.info(f"Qdrant collection '{QDRANT_COLLECTION}' created")

        # Re-index existing chunks from MongoDB into Qdrant
        existing = await db.chunks.find({}, {"_id": 0}).to_list(10000)
        if existing:
            texts = [c["text"] for c in existing]
            vectors = embed_texts(texts)
            points = [
                PointStruct(
                    id=c["id"],
                    vector=vec,
                    payload={
                        "document_id": c["document_id"],
                        "text": c["text"],
                        "metadata": c.get("metadata", {}),
                    }
                )
                for c, vec in zip(existing, vectors)
            ]
            vector_store.upsert(collection_name=QDRANT_COLLECTION, points=points)
            logger.info(f"Re-indexed {len(points)} chunks into Qdrant")
    except Exception as e:
        logger.warning(f"Qdrant setup warning: {e}")

    logger.info("Atlas AI v2.0 started (Agent Pipeline + Qdrant)")


@app.on_event("shutdown")
async def shutdown():
    client.close()
